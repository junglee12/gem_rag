# app.py
import streamlit as st
import google.generativeai as genai
import io
import logging
# Added Any for status object type hint flexibility
from typing import List, Optional, Dict, Tuple, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pypdf
import docx
import pandas as pd
from PIL import Image
import faiss
import numpy as np
import time
import datetime  # For timestamps in the log

# --- Constants ---
AVAILABLE_GEN_MODELS = [
    "gemini-2.0-flash-lite",
    "gemini-2.0-flash",
    "gemini-2.0-flash-thinking-exp-01-21",
    "gemini-2.5-pro-exp-03-25",
]
DEFAULT_GEN_MODEL = AVAILABLE_GEN_MODELS[0]
EMBEDDING_MODEL = "models/embedding-001"  # Or "models/text-embedding-004"
CHUNK_SIZE, CHUNK_OVERLAP = 1000, 150
TOP_K, HISTORY_LIMIT, API_TIMEOUT = 5, 5, 120
SUPPORTED_EXTENSIONS = ["pdf", "docx", "xlsx", "png", "jpg", "jpeg"]

# --- Logging ---
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# --- Session State Initialization ---
defaults = {
    "processed": False,
    "vector_index": None,
    "text_chunks": [],
    "messages": [],  # Main chat history (NOW ONLY user, assistant)
    "system_log": [],  # Persistent detailed log for display below chat
    "selected_model": DEFAULT_GEN_MODEL,
    "analyze_images": False,
    "api_configured": False,
    "user_api_key": None,  # API Key is preserved on clear
    "processed_filenames": [],
}
for key, value in defaults.items():
    st.session_state.setdefault(key, value)

# --- Helper Function to Add System Messages to Chat History (REMOVED FUNCTIONALITY) ---
# This function is no longer used to add messages to the main chat display.
# Kept definition in case it's needed for other future purposes, but it does nothing now.


def add_system_message(message: str):
    """NO LONGER ADDS messages to the main chat history. Use add_log_entry."""
    # st.session_state.messages.append({"role": "system", "content": message}) # <-- REMOVED
    # logger.info(f"System Message Intended for Chat (Not Added): {message}") # Optional: log if needed
    pass  # Does nothing now

# --- Helper Function to Add Messages to the Persistent System Log ---


def add_log_entry(message: str, status: Optional[Any] = None):
    """Adds a timestamped message to the persistent system log and optionally updates st.status."""
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    log_message = f"[{timestamp}] {message}"
    # Ensure log doesn't grow indefinitely (optional limit)
    # MAX_LOG_ENTRIES = 500
    # st.session_state.system_log = st.session_state.system_log[-MAX_LOG_ENTRIES:]
    st.session_state.system_log.append(log_message)
    # Also write to the status object if provided, for real-time sidebar feedback
    if status:
        # Write the message without timestamp to status for cleaner look
        status.write(message)

# --- Utility Functions ---
# Modified handle_api_error to use add_log_entry ONLY


def handle_api_error(e: Exception, context: str, model: str = "", status: Optional[Any] = None) -> None:
    """Handle API errors: log, show st.error, update status, add to persistent log."""
    error_msg = str(e).lower()
    logger.error(f"API Error in {context}: {e}")
    user_message = f"An API error occurred ({context})."

    if "api key not valid" in error_msg or ("permission_denied" in error_msg and "api key" in error_msg):
        user_message = "Invalid or expired API key. Check sidebar."
        st.session_state.api_configured = False
        st.session_state.user_api_key = None
    elif "permission denied" in error_msg:
        user_message = f"Permission denied for model '{model}'."
    elif "resource_exhausted" in error_msg:
        user_message = f"API quota exceeded or resource unavailable. Please try again later. ({context})"
    elif "deadline_exceeded" in error_msg or "timeout" in error_msg:
        user_message = f"API request timed out. Please try again. ({context})"
    elif "internal" in error_msg:
        user_message = f"Internal API error. Please try again later. ({context})"
    elif "invalid_argument" in error_msg:
        user_message = f"Invalid request sent to API. ({context})"
    else:
        user_message = f"API error: {e} ({context})."

    st.error(user_message)  # Display red error banner

    # Add detailed error to persistent log and update status
    add_log_entry(f"❌ Error: {user_message}", status=status)

    # REMOVED: add_system_message(f"Error during '{context}'. Check system log below for details.")

    # Optionally mark the status container as errored
    if status:
        status.update(state="error")


def get_model(model_name: str, status: Optional[Any] = None) -> Optional[genai.GenerativeModel]:
    """Load a generative model, handling potential errors and updating status/log."""
    if not st.session_state.api_configured:
        msg = "Cannot load model: API key not configured."
        add_log_entry(f"⚠️ {msg}", status=status)
        return None
    try:
        model = genai.GenerativeModel(model_name)
        # add_log_entry(f"Model '{model_name}' loaded.", status=status) # Optional: success log
        return model
    except Exception as e:
        # Pass status to handle_api_error, which will log and update status
        handle_api_error(
            e, f"Loading model '{model_name}'", model_name, status=status)
        return None

# --- Image Description (uses add_log_entry) ---


def describe_image(image_bytes: bytes, filename: str, model_name: str, status: Optional[Any] = None) -> Optional[str]:
    """Generate image description, updating status and log."""
    model = get_model(model_name, status=status)
    if not model:
        return None  # Error already logged by get_model/handle_api_error

    try:
        add_log_entry(f"⏳ Analyzing image: {filename}...", status=status)
        img = Image.open(io.BytesIO(image_bytes))
        prompt_parts = [
            "Describe this image in detail. If there is text, transcribe it accurately.",
            img,
        ]
        response = model.generate_content(
            prompt_parts, request_options={"timeout": API_TIMEOUT})
        desc = response.text.strip() if hasattr(
            response, 'text') and response.text else None
        if desc:
            add_log_entry(
                f"✔️ Finished analyzing image: {filename}.", status=status)
        else:
            add_log_entry(
                f"⚠️ Could not get description for image: {filename}.", status=status)
        return desc
    except Exception as e:
        # Pass status to handle_api_error, which will log and update status
        handle_api_error(
            e, f"Analyzing image '{filename}'", model_name, status=status)
        return None

# --- File Loaders (use add_log_entry) ---


def load_pdf(file, analyze_images: bool, gen_model_name: str, status: Optional[Any] = None) -> Optional[str]:
    """Extract text/images from PDF, updating status and log."""
    content = []
    file_name = file.name
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(file.getvalue()))
        num_pages = len(pdf_reader.pages)
        add_log_entry(
            f"📄 Processing PDF '{file_name}' ({num_pages} pages)...", status=status)

        for i, page in enumerate(pdf_reader.pages):
            page_num = i + 1
            # Update the main status label in sidebar for visual progress
            if status:
                status.update(
                    label=f"Processing PDF '{file_name}' - Page {page_num}/{num_pages}")
            try:
                text = page.extract_text() or ""
                if text:
                    content.append(
                        f"\n--- PDF Page {page_num} Text ---\n{text.strip()}")
                    # add_log_entry(f"   - Extracted text from page {page_num}", status=status) # Optional verbosity

                if analyze_images and st.session_state.api_configured and page.images:
                    num_images = len(page.images)
                    add_log_entry(
                        f"   - Found {num_images} images on page {page_num}. Analyzing...", status=status)
                    for j, img in enumerate(page.images):
                        img_filename_part = f"{file_name} (Page {page_num}, Img {j+1})"
                        # Pass status down to describe_image (which uses add_log_entry)
                        desc = describe_image(
                            img.data, img_filename_part, gen_model_name, status=status)
                        if desc:
                            content.append(
                                f"\n--- PDF Page {page_num} Image {j+1} Description ---\n{desc}")

            except Exception as page_e:
                logger.error(
                    f"Error processing page {page_num} of PDF '{file.name}': {page_e}")
                add_log_entry(
                    f"   - ⚠️ Error processing page {page_num}: {page_e}", status=status)
                content.append(
                    f"\n--- Error processing PDF Page {page_num} ---")

        result = "\n".join(content).strip()
        add_log_entry(f"✔️ Finished PDF: '{file_name}'.", status=status)
        return result if result else None
    except Exception as e:
        logger.error(f"Failed to process PDF '{file.name}': {e}")
        error_message = f"Failed to load PDF '{file_name}': {e}. Might be corrupted/protected."
        # Add critical error to persistent log
        add_log_entry(f"❌ {error_message}", status=status)
        # REMOVED: add_system_message(f"Error loading PDF '{file_name}'. Check system log.")
        if status:
            status.update(state="error")
        return "\n".join(content).strip() or None  # Return any partial content


def load_docx(file, status: Optional[Any] = None) -> Optional[str]:
    """Extract text from DOCX, updating status and log."""
    file_name = file.name
    try:
        add_log_entry(f"📄 Processing DOCX: {file_name}...", status=status)
        doc = docx.Document(io.BytesIO(file.getvalue()))
        full_text = []
        for p in doc.paragraphs:
            if p.text:
                full_text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                full_text.append(row_text)
        result = "\n".join(full_text).strip()
        add_log_entry(f"✔️ Finished DOCX: {file_name}.", status=status)
        return result or None
    except Exception as e:
        logger.error(f"Failed to process DOCX '{file.name}': {e}")
        error_message = f"Failed to process DOCX '{file_name}': {e}"
        add_log_entry(f"❌ {error_message}", status=status)
        # REMOVED: add_system_message(f"Error loading DOCX '{file_name}'. Check system log.")
        if status:
            status.update(state="error")
        return None


def load_excel(file, status: Optional[Any] = None) -> Optional[str]:
    """Extract text from Excel sheets, updating status and log."""
    file_name = file.name
    try:
        add_log_entry(f"📄 Processing Excel: {file_name}...", status=status)
        sheets = pd.read_excel(io.BytesIO(file.getvalue()),
                               sheet_name=None, engine='openpyxl')
        content = []
        sheet_names = list(sheets.keys())
        add_log_entry(
            f"   - Found sheets: {', '.join(sheet_names)}", status=status)
        for name, df in sheets.items():
            if not df.empty:
                sheet_content = df.to_string(index=False, na_rep='NA').strip()
                content.append(
                    f"\n--- Excel Sheet: {name} ---\n{sheet_content}")
            else:
                add_log_entry(
                    f"   - Sheet '{name}' is empty. Skipping.", status=status)
        result = "\n".join(content).strip()
        add_log_entry(f"✔️ Finished Excel: {file_name}.", status=status)
        return result or None
    except Exception as e:
        logger.error(f"Failed to process Excel '{file.name}': {e}")
        error_message = f"Failed to process Excel '{file_name}': {e}"
        add_log_entry(f"❌ {error_message}", status=status)
        # REMOVED: add_system_message(f"Error loading Excel '{file_name}'. Check system log.")
        if status:
            status.update(state="error")
        return None


def load_image(file, gen_model_name: str, status: Optional[Any] = None) -> Optional[str]:
    """Describe an image file, updating status and log."""
    file_name = file.name
    if not st.session_state.api_configured:
        msg = f"Cannot analyze image {file_name}: API key not configured."
        add_log_entry(f"⚠️ {msg}", status=status)
        return None
    # Pass status down to describe_image (which uses add_log_entry)
    desc = describe_image(file.getvalue(), file_name,
                          gen_model_name, status=status)
    if desc:
        return f"\n--- Image File: {file.name} Description ---\n{desc}"
    else:
        # Failure message handled within describe_image/handle_api_error
        return None


# Map MIME types to their respective loader functions
SUPPORTED_MIME_TYPE_LOADERS = {
    "application/pdf": load_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": load_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": load_excel,
    "image/png": load_image,
    "image/jpeg": load_image,
    "image/jpg": load_image,
}

# --- Text Processing (uses add_log_entry) ---


def split_text(text: str, status: Optional[Any] = None) -> List[str]:
    """Split text into manageable chunks, updating status and log."""
    if not text:
        return []
    add_log_entry("Splitting document text into chunks...", status=status)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        add_start_index=True,  # Helps debug chunk boundaries if needed
    )
    chunks = splitter.split_text(text)
    add_log_entry(f"Split text into {len(chunks)} chunks.", status=status)
    return chunks

# --- Embedding (Cached function, status/log updated before/after call) ---


# Spinner during actual API call
@st.cache_data(show_spinner="Embedding document chunks...")
def get_embeddings_for_chunks(_chunks_tuple: Tuple[str, ...]) -> Optional[List[List[float]]]:
    """Generate embeddings for text chunks. Input must be hashable (tuple)."""
    # Log entries are added by the caller *before* and *after* this function runs.
    chunks = list(_chunks_tuple)
    if not chunks:
        return None
    if not st.session_state.api_configured:
        return None  # Caller should log this
    try:
        logger.info(
            f"Requesting embeddings for {len(chunks)} chunks (cache may be used)...")
        response = genai.embed_content(
            model=EMBEDDING_MODEL, content=chunks, task_type="retrieval_document"
        )
        embeddings = response.get("embedding")
        if embeddings and len(embeddings) == len(chunks):
            logger.info(f"Successfully received {len(embeddings)} embeddings.")
            return embeddings
        else:
            logger.error(f"Embedding response missing or length mismatch.")
            return None  # Caller handles error logging/status update
    except Exception as e:
        logger.error(f"API error during embedding generation: {e}")
        # Cannot call handle_api_error here as it needs status. Caller must check for None return.
        return None

# --- RAG Pipeline (Errors logged, no chat messages needed here) ---


def embed_query(query: str) -> Optional[np.ndarray]:
    """Embed the user query for retrieval."""
    # No need for add_log_entry here, focus log on processing phase
    if not st.session_state.api_configured:
        add_log_entry(
            "⚠️ Cannot embed query: API key not configured.")  # Log it
        # REMOVED: add_system_message("Cannot embed query: API key not configured.")
        return None
    try:
        response = genai.embed_content(
            model=EMBEDDING_MODEL, content=query, task_type="retrieval_query"
        )
        embedding = response.get("embedding")
        if embedding:
            return np.array(embedding, dtype="float32").reshape(1, -1)
        else:
            logger.error("Embedding response for query was empty.")
            add_log_entry(
                "❌ Error: Failed to embed query (empty response).")  # Log it
            # REMOVED: add_system_message("Error: Failed to embed query.")
            return None
    except Exception as e:
        handle_api_error(e, "Embedding query")  # Logs the error
        return None


def retrieve_chunks(query_emb: np.ndarray, index: Optional[faiss.Index], chunks: List[str]) -> str:
    """Retrieve relevant chunks from the FAISS index."""
    # No need for add_log_entry here unless error
    if index is None or index.ntotal == 0 or not chunks:
        return ""
    try:
        k = min(TOP_K, index.ntotal)
        distances, indices = index.search(query_emb, k)
        valid_indices = [i for i in indices[0] if 0 <= i < len(chunks)]
        retrieved = [chunks[i] for i in valid_indices]
        if not retrieved:
            return ""
        context = "\n\n---\n\n".join(retrieved)
        # logger.info(f"Retrieved {len(retrieved)} chunks.") # Optional debug log
        return context
    except Exception as e:
        logger.error(f"Error during FAISS search: {e}")
        add_log_entry(f"❌ Error retrieving document parts: {e}")  # Log it
        # REMOVED: add_system_message(f"Error retrieving document parts: {e}")
        return ""


def format_history(history: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """Format chat history for model, filtering non-user/assistant messages."""
    max_messages = HISTORY_LIMIT * 2
    # Filter explicitly for user and assistant roles
    model_history = [msg for msg in history if msg["role"]
                     in ["user", "assistant"]]
    return model_history[-max_messages:]


def get_rag_response(query: str, index: Optional[faiss.Index], chunks: List[str], history: List[Dict[str, str]], model_name: str) -> str:
    """Generate a response using the RAG pipeline."""
    # No need for add_log_entry here, focus log on processing phase
    start_time = time.time()
    if not st.session_state.api_configured:
        add_log_entry(
            "⚠️ RAG response generation skipped: API key not configured.")
        return "API key not configured."

    context = ""
    if index is None or index.ntotal == 0 or not chunks:
        context = "No document context available."
        add_log_entry("ℹ️ Generating response without document context.")
    else:
        logger.info("Embedding user query...")
        query_emb = embed_query(query)
        if query_emb is None:
            # Error already logged by embed_query
            return "Failed to process query (embedding error)."

        logger.info("Retrieving relevant document chunks...")
        context = retrieve_chunks(query_emb, index, chunks)
        if not context:
            context = "No specific information found in documents."
            add_log_entry("ℹ️ No relevant document parts found for the query.")
        # else: logger.info("Context retrieved for RAG.") # Optional log

    formatted_history = format_history(history)  # Already filters roles
#     system_prompt = f"""Answer based *only* on the provided context and chat history. If the context or history lacks the answer, state that clearly. Do not make assumptions or use external knowledge.
# Context:
# ---
# {context}
# ---
# """
    system_prompt = f"""Act as an AI Tutor for any subject. Your primary goal is to help the user learn and understand concepts effectively.

        **Your Tutoring Approach:**
        *   **Be Patient & Encouraging:** Create a positive and supportive learning environment. Assume the user is trying their best.
        *   **Explain Clearly:** Break down complex topics into smaller, digestible steps. Use examples and analogies relevant to the subject.
        *   **Check Understanding:** Regularly ask clarifying questions to gauge the user's comprehension (e.g., "Does that make sense?", "Can you try explaining that in your own words?", "What part is confusing?"). Don't move on too quickly.
        *   **Guide, Don't Just Answer:** When asked a question, especially one that requires problem-solving, try to guide the user towards the answer themselves. Ask leading questions (Socratic method). Help them understand the *process*.
        *   **Adapt:** Adjust your explanations based on the user's responses and perceived level of understanding. If one explanation doesn't work, try another approach.
        *   **Give Feedback:** Provide constructive feedback on the user's attempts. Gently correct misunderstandings and explain *why* something is incorrect. Praise effort and correct reasoning.

        **Using Provided Context (Uploaded Documents):**
        *   If relevant context from uploaded documents is provided below ({context != "No document context available." and context != "No specific information found in documents."}), **prioritize** using that information as the basis for your explanations and examples. Refer to it explicitly if helpful (e.g., "As the provided text mentions...").
        *   If the user asks about something related to the document's topic but *not* explicitly covered in the provided context, you can use your general knowledge but clearly state that the specific detail isn't in the material provided.

        **Using General Knowledge:**
        *   If no relevant documents are provided (context is empty or clearly irrelevant), use your broad general knowledge to tutor the user on the subject they ask about.

        **Important Constraints:**
        *   **Do NOT simply give away answers**, especially for homework, assignments, or test questions. Focus on explaining the underlying concepts and problem-solving steps.
        *   Maintain a friendly, professional, and encouraging tone throughout the interaction.
        *   Stay focused on the user's learning goals for the current topic.

        ---
        **Context from Uploaded Documents (if available):**
        {context} 
        ---
        """
    logger.info(f"Generating response using model {model_name}...")
    model = get_model(model_name)  # No status needed here
    if not model:
        # Error logged by get_model/handle_api_error
        return f"Model '{model_name}' unavailable."

    model_input_history = []
    for msg in formatted_history:
        # Map to 'user' and 'model' roles expected by the API
        role = "user" if msg["role"] == "user" else "model"
        model_input_history.append({"role": role, "parts": [msg["content"]]})

    try:
        # Construct the final prompt including history, system instructions, context, and the new query
        full_prompt_parts = model_input_history + \
            [{"role": "user", "parts": [system_prompt + "\nUser Query: " + query]}]

        response = model.generate_content(
            full_prompt_parts, request_options={"timeout": API_TIMEOUT})
        logger.info(
            f"Response generated in {time.time() - start_time:.2f} seconds.")

        response_text = response.text.strip() if hasattr(
            response, 'text') and response.text else ""

        # Check for blocked response
        if not response_text and hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
            reason = response.prompt_feedback.block_reason
            add_log_entry(
                f"⚠️ Response blocked due to safety settings ({reason}). Query: '{query[:50]}...'")  # Log it
            # REMOVED: add_system_message(f"Response blocked due to safety settings ({reason}).")
            return f"Response blocked by safety settings ({reason}). Please rephrase your query or check the model's safety configuration."
        elif not response_text:
            add_log_entry(
                f"⚠️ Model returned an empty response. Query: '{query[:50]}...'")  # Log it
            return "Model did not provide an answer."
        else:
            # logger.info("Response successfully generated.") # Optional log
            return response_text

    except Exception as e:
        # handle_api_error logs the error
        handle_api_error(e, "Generating response", model_name)
        return "Error generating response. Please check the system log for details."


# --- Streamlit UI ---
st.set_page_config(page_title="Doc Q&A", layout="wide",
                   initial_sidebar_state="expanded")
st.title("💬 Document Q&A with Gemini")

# --- Sidebar ---
with st.sidebar:
    st.header("🔑 API Configuration")
    api_key_input = st.text_input(
        "Google AI API Key:", type="password", value=st.session_state.user_api_key or "",
        help="Get your key from Google AI Studio (https://aistudio.google.com/app/apikey)."
    )
    if st.button("Set API Key"):
        if api_key_input:
            try:
                # Test the key implicitly by configuring
                genai.configure(api_key=api_key_input)
                # Optionally, make a small test call like listing models if configure doesn't raise error
                # genai.list_models()
                st.session_state.user_api_key = api_key_input
                st.session_state.api_configured = True
                st.success("API Key configured!")
                # Log confirmation
                add_log_entry("🔑 API Key configured successfully.")
                # REMOVED: add_system_message("API Key configured.")
                time.sleep(1)
                st.rerun()
            except Exception as e:
                # handle_api_error logs the error and shows st.error
                handle_api_error(e, "API Configuration")
                st.session_state.api_configured = False
        else:
            st.warning("Please enter an API Key.")

    if st.session_state.api_configured:
        st.success("API Key is set.")
    else:
        st.warning("API Key not set.")
    st.divider()

    st.header("⚙️ Settings")
    st.session_state.selected_model = st.selectbox(
        "Select Model:", options=AVAILABLE_GEN_MODELS,
        index=AVAILABLE_GEN_MODELS.index(st.session_state.selected_model),
        disabled=not st.session_state.api_configured, help="Choose the generative model."
    )
    st.session_state.analyze_images = st.toggle(
        "Analyze Images in PDFs", value=st.session_state.analyze_images,
        disabled=not st.session_state.api_configured,
        help="Enable to describe images in PDFs (slower, uses more API quota)."
    )
    st.divider()

    st.header("📄 Upload & Process")
    uploaded_files = st.file_uploader(
        "Upload Documents:", type=SUPPORTED_EXTENSIONS, accept_multiple_files=True,
        disabled=not st.session_state.api_configured,
        help="Upload PDF, DOCX, XLSX, PNG, JPG, JPEG files."
    )

    # --- Process Button Logic using st.status and persistent log ---
    if st.button("Process Uploaded Files", disabled=not uploaded_files or not st.session_state.api_configured):
        new_files = [
            f for f in uploaded_files if f.name not in st.session_state.processed_filenames]
        if not new_files:
            st.info("All uploaded files have already been processed.")
        else:
            num_new_files = len(new_files)
            st.info(f"Starting processing for {num_new_files} new file(s)...")
            # Clear previous persistent log for this run? Optional, maybe better to append.
            # st.session_state.system_log = [] # Decide if you want to clear log on each process run
            add_log_entry(
                f"🚀 Starting processing for {num_new_files} new file(s)...")

            # Use st.status for real-time VISUAL feedback in the sidebar
            with st.status(f"Processing {num_new_files} file(s)...", expanded=True) as status:
                all_content = []
                newly_processed_filenames = []
                processing_successful = True  # Assume success initially

                # --- File Loading Loop ---
                for file in new_files:
                    # Update sidebar status label
                    status.update(label=f"Processing: {file.name}")
                    # Log file start
                    add_log_entry(f"⏳ Loading file: {file.name} ({file.type})")
                    loader_func = SUPPORTED_MIME_TYPE_LOADERS.get(file.type)
                    if loader_func:
                        extracted_text = None
                        try:
                            # Prepare args based on loader function needs
                            args = [file]
                            kwargs = {'status': status}
                            if loader_func in [load_pdf, load_image]:
                                args.append(st.session_state.selected_model)
                            if loader_func == load_pdf:
                                # Insert analyze_images flag at the correct position
                                args.insert(1, st.session_state.analyze_images)

                            # Loaders now use add_log_entry(..., status=status) internally
                            extracted_text = loader_func(*args, **kwargs)

                            if extracted_text:
                                all_content.append(extracted_text)
                                newly_processed_filenames.append(file.name)
                                # Don't log success here, loader function does it
                            else:
                                # Loader function should have logged the specific reason for no text
                                add_log_entry(
                                    f"ℹ️ No text content extracted or error occurred for: {file.name}", status=status)
                                # Decide if a specific loader failing means overall failure
                                # processing_successful = False # Uncomment if any file failure should stop indexing

                        except Exception as load_e:
                            logger.error(
                                f"Unhandled error loading file {file.name}: {load_e}")
                            error_msg = f"Critical error loading {file.name}. Skipping. Error: {load_e}"
                            # Log & update status
                            add_log_entry(f"❌ {error_msg}", status=status)
                            # REMOVED: add_system_message(error_msg)
                            processing_successful = False  # Critical error stops processing
                    else:
                        add_log_entry(
                            f"⚠️ Unsupported file type: {file.name} ({file.type}). Skipping.", status=status)

                # --- Embedding Stage ---
                if all_content and processing_successful:
                    status.update(label="Combining text and splitting...")
                    add_log_entry("Combining extracted text...", status=status)
                    combined_text = "\n\n".join(all_content)
                    # split_text uses add_log_entry(..., status=status)
                    new_chunks = split_text(combined_text, status=status)

                    if new_chunks:
                        status.update(
                            label=f"Embedding {len(new_chunks)} text chunks...")
                        add_log_entry(
                            f"Generating embeddings for {len(new_chunks)} chunks using '{EMBEDDING_MODEL}'...", status=status)
                        # Call cached function (shows its own spinner)
                        new_embeddings = None
                        embedding_error = None
                        try:
                            new_embeddings = get_embeddings_for_chunks(
                                tuple(new_chunks))
                        except Exception as embed_e:
                            # Catch errors from the embedding call itself (like API key issues detected late)
                            embedding_error = embed_e
                            logger.error(
                                f"Error during get_embeddings_for_chunks call: {embed_e}")
                            # Log via handler
                            handle_api_error(
                                embed_e, "Embedding document chunks", status=status)

                        if new_embeddings:
                            add_log_entry(
                                "✔️ Embeddings generated successfully.", status=status)
                            embeddings_np = np.array(
                                new_embeddings, dtype="float32")
                            if embeddings_np.ndim == 2 and embeddings_np.shape[0] == len(new_chunks) and embeddings_np.shape[1] > 0:
                                embedding_dim = embeddings_np.shape[1]
                                if st.session_state.vector_index is None:
                                    add_log_entry(
                                        f"Creating new vector index (dim: {embedding_dim}).", status=status)
                                    st.session_state.vector_index = faiss.IndexFlatL2(
                                        embedding_dim)
                                else:
                                    add_log_entry(
                                        f"Using existing vector index (dim: {st.session_state.vector_index.d}).", status=status)

                                if st.session_state.vector_index.d == embedding_dim:
                                    add_log_entry(
                                        f"Adding {embeddings_np.shape[0]} vectors to index.", status=status)
                                    st.session_state.vector_index.add(
                                        embeddings_np)
                                    st.session_state.text_chunks.extend(
                                        new_chunks)
                                    st.session_state.processed_filenames.extend(
                                        newly_processed_filenames)
                                    st.session_state.processed = True
                                    add_log_entry(
                                        f"✔️ Index updated. Total chunks: {len(st.session_state.text_chunks)}, Index size: {st.session_state.vector_index.ntotal}", status=status)
                                else:
                                    error_msg = f"Embedding dimensions mismatch (Index: {st.session_state.vector_index.d}, New: {embedding_dim}). Cannot add new vectors."
                                    add_log_entry(
                                        f"❌ {error_msg}", status=status)
                                    # REMOVED: add_system_message(f"Error: {error_msg}")
                                    processing_successful = False
                            else:
                                error_msg = "Invalid embeddings received (shape mismatch or empty)."
                                add_log_entry(f"❌ {error_msg}", status=status)
                                # REMOVED: add_system_message(f"Error: {error_msg}")
                                processing_successful = False
                        else:
                            # Error occurred during embedding generation (API error or empty response)
                            if not embedding_error:  # If handle_api_error wasn't already called
                                error_msg = "Embedding process failed (no embeddings returned)."
                                add_log_entry(f"❌ {error_msg}", status=status)
                            # Check if a specific API error was logged by handle_api_error
                            # api_error_logged = any("Embedding document chunks" in e_msg for e_msg in st.session_state.system_log) # Check log now
                            # if not api_error_logged: add_log_entry(f"Error: {error_msg} Check log.") # Already logged above
                            processing_successful = False
                    else:
                        add_log_entry(
                            "⚠️ Content extracted, but no text chunks created after splitting (check content/splitter).", status=status)
                elif not all_content and new_files and processing_successful:
                    add_log_entry(
                        "⚠️ No text content extracted from any of the new files.", status=status)
                    processing_successful = False  # Nothing was added
                elif not processing_successful:
                    add_log_entry(
                        "⚠️ Processing halted due to earlier errors. No embedding attempted.", status=status)

                # --- Update Final Status Container State ---
                if processing_successful and st.session_state.processed:
                    final_message = f"Successfully processed {len(newly_processed_filenames)} file(s)."
                    status.update(label=final_message,
                                  state="complete", expanded=False)
                    # Log final summary
                    add_log_entry(f"✅ {final_message} Ready for questions.")
                    # REMOVED: add_system_message(f"{final_message} Ready for questions.")
                else:  # Errors occurred or no content processed
                    final_message = f"Processing finished with errors or no new data added."
                    # Ensure status reflects error if processing_successful is False
                    final_state = "error" if not processing_successful else "warning"
                    status.update(label=final_message,
                                  state=final_state, expanded=True)
                    # Log final summary
                    add_log_entry(f"⚠️ {final_message} Check log for details.")
                    # REMOVED: add_system_message(f"{final_message} Check system log below for details.")

            # --- End of st.status block ---
            # Rerun to display the populated system_log in the main area and update processed files list
            st.rerun()

    st.divider()
    st.subheader("Processed Files:")
    if st.session_state.processed_filenames:
        with st.expander(f"{len(st.session_state.processed_filenames)} file(s) processed", expanded=False):
            for filename in st.session_state.processed_filenames:
                st.markdown(f"- `{filename}`")
    else:
        st.caption("No files processed yet.")
    st.divider()

    # st.subheader("System Processing Log")
    # # Set height for scrollable log area
    # log_container = st.container(height=200)
    # with log_container:
    #     if st.session_state.system_log:
    #         # Display log entries chronologically (newest at the bottom implicitly)
    #         log_html = "<div style='font-family: monospace; font-size: small;'>"
    #         log_html += "<br>".join(st.session_state.system_log)
    #         log_html += "</div>"
    #         # Use markdown with unsafe_allow_html for better control if needed, or just iterate
    #         # st.markdown(log_html, unsafe_allow_html=True)
    #         # Simpler approach: iterate and use markdown with backticks
    #         for log_entry in st.session_state.system_log:
    #             st.markdown(f"`{log_entry}`")

    #     else:
    #         # Placeholder text
    #         st.caption("Processing log will appear here...")
    # st.divider()  # Optional: Add another divider before Clear button

    # --- Clear Chat History Button ---
    st.header("🗑️ Clear Chat")  # Changed header slightly for clarity
    # CHANGE Button Label and Help Text
    if st.button("Clear Chat History & Log", help="Resets only the chat history and system log. Processed documents remain."):
        # CHANGE keys_to_reset to ONLY include 'messages' and 'system_log'
        keys_to_reset = ["messages", "system_log"]

        # Keep API key preservation logic (it doesn't hurt)
        api_key = st.session_state.user_api_key
        api_conf = st.session_state.api_configured
        sel_model = st.session_state.selected_model

        # Reset only the specified keys (messages and system_log) to their defaults
        for key in keys_to_reset:
            if key in defaults:
                # Reset to default (e.g., [])
                st.session_state[key] = defaults[key]
            elif hasattr(st.session_state, key):
                # If somehow it's not in defaults but exists, remove it (less likely for these keys)
                del st.session_state[key]

        # Restore preserved settings
        st.session_state.user_api_key = api_key
        st.session_state.api_configured = api_conf
        st.session_state.selected_model = sel_model

        # REMOVE the cache clearing - we want to keep embeddings if data is kept
        # st.cache_data.clear() # <--- COMMENT OUT or DELETE this line

        # CHANGE Success message
        st.success("Chat history and system log cleared.")
        # CHANGE Log entry
        # Log confirmation
        add_log_entry("🧹 Chat history and system log cleared.")
        time.sleep(1)
        st.rerun()  # Rerun to update the display

        # --- Add a divider between the clear buttons ---

    # --- Clear Processed Data Button ---
    # You might want a more general header now, like "🗑️ Manage Session Data"
    # Or just keep separate headers/buttons. Let's add a new button under the same header for now.
    # st.header("🗑️ Clear Processed Data") # Optional: Add another header if preferred

    if st.button("Clear Processed Data", help="Removes loaded documents, index, and cached embeddings. Keeps chat history and API key."):
        # Define keys specific to processed data
        keys_to_reset_data = [
            "processed",
            "vector_index",
            "text_chunks",
            "processed_filenames"
        ]

        # Keep API key preservation logic (optional but safe)
        api_key = st.session_state.user_api_key
        api_conf = st.session_state.api_configured
        sel_model = st.session_state.selected_model

        # Reset only the specified data keys to their defaults
        for key in keys_to_reset_data:
            if key in defaults:
                st.session_state[key] = defaults[key]  # Reset to default
            elif hasattr(st.session_state, key):
                # If somehow it's not in defaults but exists, remove it
                del st.session_state[key]

        # Restore preserved settings
        st.session_state.user_api_key = api_key
        st.session_state.api_configured = api_conf
        st.session_state.selected_model = sel_model

        # IMPORTANT: Clear the embedding cache when clearing processed data
        st.cache_data.clear()

        # Update Success message
        st.success("Processed document data and cache cleared.")
        # Update Log entry
        # Log confirmation
        add_log_entry("🧹 Processed document data and embedding cache cleared.")
        time.sleep(1)
        st.rerun()  # Rerun to update the display (e.g., processed files list)
    st.divider()

# --- Main Area ---

# 1. Handle input FIRST and append to state
# DEFINE PLACEHOLDER AND ENABLEMENT *BEFORE* USING THEM
chat_enabled = st.session_state.api_configured and (st.session_state.processed or st.session_state.processed_filenames)
prompt_placeholder = "Ask a question about the documents..." if (st.session_state.processed or st.session_state.processed_filenames) else "Process documents first..."
# NOW CALL CHAT_INPUT
prompt = st.chat_input(prompt_placeholder, disabled=not chat_enabled)
if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    # No rerun needed here

# 2. Define chat display area and display history
st.subheader("Chat")
chat_container = st.container() # Use dynamic height if preferred
with chat_container:
    if not st.session_state.api_configured:
        st.warning("👈 Configure API Key in sidebar to enable processing and chat.")
    elif not st.session_state.processed and not st.session_state.processed_filenames:
        st.info("👈 Upload & Process documents in sidebar to begin chat.")

    # Display all messages currently in state. This now includes the user message added above.
    for message in st.session_state.messages:
         if message["role"] in ["user", "assistant"]:
             with st.chat_message(message["role"]):
                 st.markdown(message["content"])

    # Create a placeholder *inside* the container *after* the loop.
    # This is where the assistant's response will appear for the CURRENT turn.
    assistant_response_area = st.empty()

# 3. Generate and display assistant response IN THE PLACEHOLDER
# Check if the last message is 'user', indicating a response is needed for this turn.
if chat_enabled and st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    # Use the placeholder context manager to manage the spinner and final response display
    with assistant_response_area.container():
         with st.chat_message("assistant"): # Create the assistant message bubble structure
             with st.spinner("Thinking..."): # Show spinner inside the bubble
                 # Generate the actual response
                 response_content = get_rag_response(
                     st.session_state.messages[-1]["content"], # Last message is the user query
                     st.session_state.vector_index,
                     st.session_state.text_chunks,
                     st.session_state.messages[:-1], # History *before* the user query
                     st.session_state.selected_model
                 )
                 # Display the response, replacing the spinner within the placeholder's bubble
                 st.markdown(response_content)

    # Append the generated response to the state *after* it has been displayed in the placeholder.
    # This ensures the state reflects the conversation history for the *next* turn.
    st.session_state.messages.append({"role": "assistant", "content": response_content})

    # We should NOT need a rerun here. The display happened via the placeholder.
    # The state is updated for the next interaction.

# 4. Chat Input widget definition (can be placed earlier or later in code, Streamlit handles layout)
# We already defined it at the top to capture the prompt value early.